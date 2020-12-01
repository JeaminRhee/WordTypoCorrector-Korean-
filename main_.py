from docx import Document
import docx2txt
import os
from tkinter import *


def show_typos(misspelled_dict):
    def got_it():
        root.destroy()
    root = Tk()
    root.title("Typo&Correction Window")
    scroll = Scrollbar(root)
    scroll.pack(side=RIGHT,fill=Y)

    T = Text(root, height=25, width=47, yscrollcommand=scroll.set)
    T.insert(END, "TYPO\t\t  CORRECT\n")
    T.pack()
    B = Button(root, text="Got it", width=6, command=got_it)
    B.pack()
    scroll.config(command=T.yview)

    for typo in misspelled_dict.keys():
        T.insert(END, typo+"\t\t  "+misspelled_dict.get(typo)+"\n")

    root.mainloop()


def typolistread():
    f = open('TypoCollection.csv', 'r')
    content = f.readlines()

    typo_dict = {}

    for line in content:
        items = line.split(',')
        key, values = items[0], items[1]
        values = values[:-1]
        typo_dict[key] = values
    f.close()
    return typo_dict


def correcting_process(typo_dict, my_text):
    misspelled_dict = {}
    for correction in typo_dict.keys():
            if correction in my_text:
                my_text = my_text.replace(correction, typo_dict.get(correction))
                misspelled_dict.update({correction:typo_dict.get(correction)})

    return my_text, misspelled_dict


def docx_process(savefilename, typo_dict):
    savefilename = savefilename + ".docx"
    # my_text is a string that has all text in the initial docx file
    # Make sure to have only one file per extension so that TypoDetector can automatically correct.
    for files in os.listdir("."):
        if files.endswith(".docx"):
            my_text = docx2txt.process(files)
            break
    doc = Document()
    para = doc.add_paragraph()
    my_text, misspelled_dict = correcting_process(typo_dict, my_text)
    run = para.add_run(my_text)
    doc.save(savefilename)
    return misspelled_dict


def txt_process(savefilename, typo_dict):
    savefilename = savefilename + ".txt"
    # my_text is a string that has all text in the initial txt file
    # Make sure to have only one file per extension so that TypoDetector can automatically correct.
    for files in os.listdir("."):
        if files.endswith(".txt"):
            f = open(files,'r')
            my_text = f.read()
            break
    of = open(savefilename,"w+")
    my_text, misspelled_dict = correcting_process(typo_dict, my_text)
    of.write(my_text)
    of.close()
    return misspelled_dict
